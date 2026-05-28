"""
Verification Report Ingestion Script
=====================================

`verification-reports/` 비공개 폴더에 위치한 docx 검증 리포트를
- 익명화 규칙 적용
- skills/{slug}/examples/{report_id}.md (마크다운 발췌)
- public/sample-data/{slug}--{report_id}.json (구조화 JSON)
- data/verification-reports.json (카드 상세에 노출될 인덱스)

로 변환한다. 원본 docx는 절대 커밋되지 않으며 (.gitignore),
생성물만 커밋되어 Vercel 자동 재배포에 반영된다.

Usage:
    python scripts/ingest_verification.py
    또는 npm run ingest

Requirements:
    pip install python-docx
"""

from __future__ import annotations

import io
import json
import sys
from pathlib import Path
from typing import Any

# Windows 콘솔 cp949 → utf-8 강제 (한글·이모지 출력 안전)
if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
if hasattr(sys.stderr, "buffer"):
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")

try:
    from docx import Document
except ImportError:  # pragma: no cover
    print("ERROR: python-docx 미설치. 'pip install python-docx' 실행 후 다시 시도하세요.")
    sys.exit(1)


ROOT = Path(__file__).resolve().parent.parent
REPORTS = ROOT / "verification-reports"
PUBLIC_SAMPLE = ROOT / "public" / "sample-data"
SKILLS = ROOT / "skills"
INDEX_PATH = ROOT / "data" / "verification-reports.json"


def read_meta(docx_path: Path) -> dict | None:
    meta_path = docx_path.with_suffix(".meta.json")
    if not meta_path.exists():
        print(f"  ⚠ meta 사이드카 없음: {meta_path.name}, 건너뜀.")
        return None
    return json.loads(meta_path.read_text(encoding="utf-8"))


def anonymize(text: str, rules: dict | None) -> str:
    if not rules or not rules.get("replace"):
        return text
    out = text
    for k, v in rules["replace"].items():
        out = out.replace(k, v)
    return out


def extract(docx_path: Path) -> dict[str, Any]:
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for i, t in enumerate(doc.tables):
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        tables.append({"index": i, "rows": rows})
    return {"paragraphs": paragraphs, "tables": tables}


def write_md(meta: dict, content: dict, anon) -> Path:
    parts = [
        f"# {meta['title']}",
        "",
        f"> 자동 추출본 — `{meta['report_id']}` ({meta['date']})  ",
        f"> 익명화 적용. 원본은 사내 보관 (verification-reports/).",
        "",
        "## 본문",
        "",
    ]
    parts += [anon(p) for p in content["paragraphs"]]
    parts.append("")
    parts.append("## 표")
    parts.append("")
    for t in content["tables"]:
        if not t["rows"]:
            continue
        parts.append(f"### Table {t['index'] + 1}")
        parts.append("")
        header = [anon(c) for c in t["rows"][0]]
        parts.append("| " + " | ".join(header) + " |")
        parts.append("| " + " | ".join(["---"] * len(header)) + " |")
        for r in t["rows"][1:]:
            cells = [anon(c) for c in r]
            parts.append("| " + " | ".join(cells) + " |")
        parts.append("")
    md = "\n".join(parts)

    out = SKILLS / meta["card_slug"] / "examples" / f"{meta['report_id']}.md"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(md, encoding="utf-8")
    return out


def write_json(meta: dict, content: dict, anon) -> Path:
    sample = {
        "report_id": meta["report_id"],
        "card": meta["card_slug"],
        "title": meta["title"],
        "date": meta["date"],
        "tables": [
            {
                "index": t["index"],
                "rows": [[anon(c) for c in r] for r in t["rows"]],
            }
            for t in content["tables"]
        ],
        "source": "auto-extracted from verification report (anonymized)",
        "labels": meta.get("labels", {}),
    }
    out = PUBLIC_SAMPLE / f"{meta['card_slug']}--{meta['report_id']}.json"
    out.write_text(json.dumps(sample, ensure_ascii=False, indent=2), encoding="utf-8")
    return out


def update_index(records: list[dict]) -> None:
    """카드 슬러그별 검증 리포트 인덱스를 갱신한다."""
    INDEX_PATH.parent.mkdir(parents=True, exist_ok=True)
    by_card: dict[str, list[dict]] = {}
    for rec in records:
        slug = rec["card_slug"]
        entry = {
            "report_id": rec["report_id"],
            "title": rec["title"],
            "date": rec["date"],
            "sample_json_url": f"/sample-data/{slug}--{rec['report_id']}.json",
            "example_md_url": f"https://github.com/hyeonchoi-cpu/listeningmind-playbook/blob/main/skills/{slug}/examples/{rec['report_id']}.md",
        }
        by_card.setdefault(slug, []).append(entry)
    # 일관성을 위해 카드별 정렬
    for slug in by_card:
        by_card[slug].sort(key=lambda r: (r["date"], r["report_id"]))
    INDEX_PATH.write_text(json.dumps(by_card, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\n  ✓ 인덱스 갱신: {INDEX_PATH.relative_to(ROOT)} ({sum(len(v) for v in by_card.values())} reports)")


def main() -> None:
    if not REPORTS.exists():
        REPORTS.mkdir(parents=True)
        print(f"폴더 생성: {REPORTS.relative_to(ROOT)}")
        print("이 폴더에 .docx + .meta.json 사이드카 쌍을 두고 다시 실행하세요.")
        return

    docx_files = sorted(REPORTS.glob("*.docx"))
    if not docx_files:
        print(f"docx 파일 없음: {REPORTS.relative_to(ROOT)}")
        return

    print(f"발견된 검증 리포트: {len(docx_files)}건\n")

    records: list[dict] = []
    for docx_path in docx_files:
        meta = read_meta(docx_path)
        if not meta:
            continue

        # 필수 필드 검증
        for k in ("report_id", "card_slug", "title", "date"):
            if k not in meta:
                print(f"  ⚠ {docx_path.name}: meta 필수 필드 누락 '{k}', 건너뜀.")
                meta = None
                break
        if meta is None:
            continue

        print(f"[Ingest] {docx_path.name}")
        print(f"  → card: {meta['card_slug']} / report_id: {meta['report_id']}")

        content = extract(docx_path)
        def anon(t, m=meta):
            return anonymize(t, m.get("anonymize", {}))

        pub = meta.get("publish", {"as_example_md": True, "as_sample_json": True})
        if pub.get("as_example_md", True):
            out = write_md(meta, content, anon)
            print(f"  ✓ {out.relative_to(ROOT)}")
        if pub.get("as_sample_json", True):
            out = write_json(meta, content, anon)
            print(f"  ✓ {out.relative_to(ROOT)}")

        records.append(meta)
        print()

    if records:
        update_index(records)
        print("\n완료. 다음 단계:")
        print("  git add data/verification-reports.json public/sample-data/ skills/")
        print("  git commit -m 'docs: add verification reports'")
        print("  git push  # Vercel 자동 재배포")


if __name__ == "__main__":
    main()
